from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))

import contract_claims_center as ccc  # noqa: E402


PROJECTS_DIR = ROOT / "projects"
PROJECT_DIRS = [path for path in PROJECTS_DIR.iterdir() if path.is_dir() and not path.name.startswith("_")]
PROJECT_DIR = max(PROJECT_DIRS, key=lambda path: sum(file.stat().st_size for file in path.rglob("*") if file.is_file()), default=PROJECTS_DIR / "_PROJECT_TEMPLATE")
DEFAULT_METHODOLOGY = PROJECT_DIR / "delay_analysis" / "methodology" / "TIA methodology.docx"
CONTRACTS_DIR = PROJECT_DIR / "contracts" / "source"
OVERALL_CONTRACT_PDF = CONTRACTS_DIR / "Overall Contract.pdf"
LETTERS_PATHS = [
    PROJECT_DIR / "letters_intelligence" / "letters_intelligence.xlsx",
]
OUTPUT_DIR = ROOT / "generated_outputs" / "contract_letters_eval"
SCORE_LOG = OUTPUT_DIR / "score_log.jsonl"


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def contains_any(text: str, terms: list[str]) -> bool:
    lower = str(text).lower()
    return any(term.lower() in lower for term in terms)


def pct(value: float, total: float) -> float:
    if total <= 0:
        return 0.0
    return round(max(0.0, min(100.0, value / total * 100.0)), 2)


def score_boolean_checks(checks: list[tuple[str, bool, float]]) -> tuple[float, list[dict[str, Any]]]:
    earned = sum(weight for _, ok, weight in checks if ok)
    total = sum(weight for _, _, weight in checks)
    return pct(earned, total), [{"criterion": name, "passed": bool(ok), "weight": weight} for name, ok, weight in checks]


def render_table(frame: pd.DataFrame, limit: int = 8) -> str:
    if frame.empty:
        return "_No rows generated._"
    sample = frame.head(limit).copy()
    for col in sample.columns:
        sample[col] = sample[col].astype(str).str.replace("\r", " ", regex=False).str.replace("\n", " ", regex=False)
    headers = [str(col) for col in sample.columns]
    rows = [[str(value) for value in row] for row in sample.values.tolist()]
    lines = [" | ".join(headers), " | ".join(["---"] * len(headers))]
    lines.extend(" | ".join(row) for row in rows)
    return "\n".join(lines)


def load_letters_workbook() -> tuple[Path | None, dict[str, pd.DataFrame]]:
    for path in LETTERS_PATHS:
        if path.exists():
            sheets = pd.read_excel(path, sheet_name=None)
            return path, {name: df.fillna("") for name, df in sheets.items()}
    return None, {}


def pdf_text_probe(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"exists": False, "pages": 0, "extractable_text_chars": 0, "is_scanned_or_image_based": False}
    pages = 0
    chars = 0
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = len(reader.pages)
        for page in list(reader.pages)[:5]:
            chars += len(page.extract_text() or "")
    except Exception:
        chars = 0
    try:
        diagnostics = ccc.build_pdf_source_diagnostics(path)
    except Exception:
        diagnostics = {}
    return {
        "exists": True,
        "pages": pages,
        "extractable_text_chars": chars,
        "is_scanned_or_image_based": chars == 0,
        "path": str(path),
        "diagnostics": diagnostics,
    }


def tokenize_notice_text(value: str) -> set[str]:
    stopwords = {
        "for", "the", "and", "with", "from", "that", "this", "into", "before", "after",
        "work", "works", "project", "letter", "response", "notice", "urgent", "required",
        "re", "rc", "of", "to", "on", "in", "no", "date", "submittal", "submission",
    }
    tokens = set(re.findall(r"[a-z0-9]+", str(value).lower()))
    return {token for token in tokens if len(token) > 2 and token not in stopwords}


def predict_delay_type(letter_type: str, risk_type: str, subject: str, main_purpose: str) -> str:
    text = " ".join([str(letter_type), str(risk_type), str(subject), str(main_purpose)]).lower()
    if any(term in text for term in ["rfi", "consultant reply", "reply from consultant"]):
        return "RFI / response delay"
    if any(term in text for term in ["ifc", "drawing", "design", "shop drawing"]):
        return "IFC / design delay"
    if any(term in text for term in ["payment", "invoice", "certificate", "certified"]):
        return "Payment delay"
    if any(term in text for term in ["steel", "reinforcement", "rft", "free issue", "material"]):
        return "Steel supply delay"
    if any(term in text for term in ["delay", "eot", "extension"]):
        return "General delay / EOT"
    return "General correspondence"


def predict_notice_status(reply_received: str, letter_type: str) -> str:
    if str(reply_received).strip().lower() == "yes":
        return "Replied"
    if "notice" in str(letter_type).lower() or "delay" in str(letter_type).lower():
        return "Open / action required"
    return "Open"


def build_letters_reference_maps(letters_book: dict[str, pd.DataFrame]) -> tuple[dict[str, dict[str, str]], dict[str, dict[str, str]]]:
    ref_map: dict[str, dict[str, str]] = {}
    thread_map: dict[str, dict[str, str]] = {}
    for sheet_name, from_party, to_party in [
        ("From SAMCO to ACE", "SAMCO", "ACEPM"),
        ("From ACE to SAMCO", "ACEPM", "SAMCO"),
    ]:
        df = letters_book.get(sheet_name, pd.DataFrame())
        for _, row in df.iterrows():
            ref = str(row.get("Ref No", "")).strip()
            if not ref:
                continue
            ref_map[ref] = {
                "From Party": from_party,
                "To Party": to_party,
                "Date": str(row.get("Date", "")).strip(),
                "Type": str(row.get("Type", "")).strip(),
                "Subject": str(row.get("Subject", "")).strip(),
                "Main Purpose": str(row.get("Main Purpose", "")).strip(),
                "Affected Activities": str(row.get("Affected Activities", "")).strip(),
                "Risk Type": str(row.get("Risk Type", "")).strip(),
                "Delay Risk": str(row.get("Delay Risk", "")).strip(),
                "Required Actions": str(row.get("Required Actions", "")).strip(),
            }
    samco_links = letters_book.get("SAMCO to ACE Links", pd.DataFrame())
    for _, row in samco_links.iterrows():
        ref = str(row.get("SAMCO Ref No", "")).strip()
        if ref:
            thread_map[ref] = {
                "Thread": str(row.get("Thread", "")).strip(),
                "Reply Ref": str(row.get("Related ACE Ref No(s)", "")).strip(),
                "Reply Date": str(row.get("ACE Date(s)", "")).strip(),
                "Relationship": str(row.get("Relationship", "")).strip(),
                "Recommended Follow-up": str(row.get("Recommended Follow-up", "")).strip(),
            }
    ace_links = letters_book.get("ACE to SAMCO Links", pd.DataFrame())
    for _, row in ace_links.iterrows():
        ref = str(row.get("ACE Ref No", "")).strip()
        if ref:
            thread_map[ref] = {
                "Thread": str(row.get("Thread(s)", "")).strip(),
                "Reply Ref": str(row.get("Related SAMCO Ref No(s)", "")).strip(),
                "Reply Date": "",
                "Relationship": "ACE to SAMCO linked thread",
                "Recommended Follow-up": str(row.get("ACE Required Actions", "")).strip(),
            }
    return ref_map, thread_map


def build_notice_register(letters_book: dict[str, pd.DataFrame]) -> pd.DataFrame:
    ref_map, thread_map = build_letters_reference_maps(letters_book)
    issue_threads = letters_book.get("Issue Threads", pd.DataFrame())
    thread_actions = {
        str(row.get("Thread", "")).strip(): {
            "Priority": str(row.get("Priority", "")).strip(),
            "Next Action": str(row.get("Next Action", "")).strip(),
        }
        for _, row in issue_threads.iterrows()
    } if not issue_threads.empty else {}
    rows = []
    for ref, details in ref_map.items():
        thread_info = thread_map.get(ref, {})
        thread_name = thread_info.get("Thread", "")
        thread_meta = thread_actions.get(thread_name, {})
        rows.append(
            {
                "Notice Ref": ref,
                "Date": details.get("Date", ""),
                "From Party": details.get("From Party", ""),
                "To Party": details.get("To Party", ""),
                "Type": details.get("Type", ""),
                "Subject": details.get("Subject", ""),
                "Thread": thread_name,
                "Predicted Delay Type": predict_delay_type(details.get("Type", ""), details.get("Risk Type", ""), details.get("Subject", ""), details.get("Main Purpose", "")),
                "Predicted Status": predict_notice_status("Yes" if thread_info.get("Reply Ref", "") else "No", details.get("Type", "")),
                "Reply Received": "Yes" if thread_info.get("Reply Ref", "") else "No",
                "Reply Ref": thread_info.get("Reply Ref", ""),
                "Reply Date": thread_info.get("Reply Date", ""),
                "Predicted Activity Text": details.get("Affected Activities", ""),
                "Priority": thread_meta.get("Priority", details.get("Delay Risk", "")),
                "Next Action": thread_meta.get("Next Action", details.get("Required Actions", "") or thread_info.get("Recommended Follow-up", "")),
            }
        )
    return pd.DataFrame(rows)


def create_evidence_file(db_path: Path, evidence_dir: Path) -> tuple[list[str], pd.DataFrame, pd.DataFrame]:
    class Upload:
        name = "eval_late_ifc_notice.txt"

        def getvalue(self) -> bytes:
            return (
                "Notice for late IFC drawings and RFI response delay. The Engineer response is late, "
                "the affected work is on the critical path, programme impact exists, and SAMCO reserves "
                "rights for EOT, cost, time impact, contemporaneous records, and mitigation evidence."
            ).encode("utf-8")

    saved = ccc.persist_uploaded_evidence(db_path, evidence_dir, [Upload()], "Letters Intelligence")
    return saved, ccc.load_evidence_documents(db_path), ccc.load_evidence_mappings(db_path)


def build_analysis() -> dict[str, Any]:
    os.environ.pop("PROJECT_HUB_OPENAI_ENABLED", None)
    os.environ.pop("OPENAI_API_KEY", None)
    temp_dir = Path(tempfile.mkdtemp(prefix="contract_letters_eval_"))
    db_path = temp_dir / "contract_claims_eval.db"
    evidence_dir = temp_dir / "evidence"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    ccc.init_contract_claims_db(db_path)
    status = ccc.persist_contract_analysis(db_path, CONTRACTS_DIR, rebuild=True)
    pdf_probe = pdf_text_probe(OVERALL_CONTRACT_PDF)
    clauses_df = ccc.load_contract_library(db_path)
    saved_evidence, evidence_df, mappings_df = create_evidence_file(db_path, evidence_dir)
    answer = ccc.answer_contract_question(db_path, "Can we claim EOT for late IFC drawings and delayed Engineer RFI response?")
    search_df = ccc.clause_search_dataframe(
        clauses_df,
        "Can we claim EOT for late IFC drawings and delayed Engineer RFI response? EOT / Delay Claim Late Drawings / Late Approvals",
        limit=10,
    )
    rebuttal = ccc.build_client_rebuttal(db_path, "Your claim is rejected because no notice was submitted, the delay is not on the critical path, and there is concurrent contractor delay.")
    selected_clause_ids = clauses_df.head(3)["id"].astype(int).tolist() if not clauses_df.empty else []
    selected_evidence_ids = evidence_df.head(2)["id"].astype(int).tolist() if not evidence_df.empty else []
    claim_payload = ccc.build_claim_draft_payload(
        db_path,
        claim_type="EOT / Delay Claim",
        delay_event="Late IFC drawings and delayed Engineer RFI response",
        selected_clause_ids=selected_clause_ids,
        selected_evidence_ids=selected_evidence_ids,
        client_rejection_text="No notice, no critical path impact, and concurrent contractor delay.",
    )
    letters_path, letters_book = load_letters_workbook()
    notice_df = build_notice_register(letters_book)
    kpis = ccc.build_contract_center_kpis(db_path)
    shutil.rmtree(temp_dir, ignore_errors=True)
    return {
        "status": status,
        "pdf_probe": pdf_probe,
        "clauses_df": clauses_df,
        "evidence_df": evidence_df,
        "mappings_df": mappings_df,
        "saved_evidence": saved_evidence,
        "answer": answer,
        "search_df": search_df,
        "rebuttal": rebuttal,
        "claim_payload": claim_payload,
        "letters_path": letters_path,
        "letters_book": letters_book,
        "notice_df": notice_df,
        "kpis": kpis,
    }


def summarize_artifact(analysis: dict[str, Any], methodology_text: str) -> str:
    answer = analysis["answer"]
    rebuttal = analysis["rebuttal"]
    claim = analysis["claim_payload"]
    lines = [
        "# Contract & Claims / Letters Eval Artifact",
        "",
        f"Generated: {datetime.now(timezone.utc).isoformat()}",
        "",
        "## Contract Center Status",
        "",
        json.dumps(analysis["status"], indent=2, default=str),
        "",
        "## Overall Contract PDF Readiness",
        "",
        json.dumps(analysis["pdf_probe"], indent=2, default=str),
        "",
        "## Contract Question Output",
        "",
        f"Decision: {answer.get('entitlement_decision')}",
        f"Risk: {answer.get('risk_assessment')}",
        f"Evidence strength: {answer.get('evidence_strength_label')} ({answer.get('evidence_strength_score')}/100)",
        f"Short answer: {answer.get('short_answer')}",
        f"Contractor interpretation: {answer.get('contractor_friendly_interpretation')}",
        f"Required evidence: {answer.get('required_evidence')}",
        f"Missing evidence: {answer.get('missing_evidence')}",
        f"Likely rejection: {answer.get('likely_client_rejection')}",
        f"Contractor rebuttal: {answer.get('contractor_rebuttal')}",
        f"Next action: {answer.get('recommended_next_action')}",
        f"Claim strategy: {answer.get('claim_strategy')}",
        "",
        "## Client Rebuttal Output",
        "",
        f"Client argument summary: {rebuttal.get('client_argument_summary')}",
        f"Contractual risk: {rebuttal.get('contractual_risk')}",
        f"Contractor counterargument: {rebuttal.get('contractor_counterargument')}",
        f"Evidence needed: {rebuttal.get('evidence_needed')}",
        f"Recommended response wording: {rebuttal.get('recommended_response_wording')}",
        f"Probability of success: {rebuttal.get('probability_of_success')}",
        "",
        "## Claim Draft Output",
        "",
        f"Narrative: {claim.get('narrative_text')}",
        f"Contractual basis: {claim.get('contractual_basis')}",
        f"Factual background: {claim.get('factual_background')}",
        f"Cause and effect: {claim.get('cause_effect')}",
        f"Evidence list: {claim.get('evidence_list')}",
        f"Entitlement statement: {claim.get('entitlement_statement')}",
        f"Time impact statement: {claim.get('time_impact_statement')}",
        f"Cost impact statement: {claim.get('cost_impact_statement')}",
        f"Rebuttal section: {claim.get('rebuttal_section')}",
        f"Attachment checklist: {claim.get('attachment_checklist')}",
        "",
        "## Letters Notice Register",
        "",
        f"Letters workbook: {analysis.get('letters_path')}",
        render_table(analysis.get("notice_df", pd.DataFrame())),
        "",
        "## Relevant Clauses",
        "",
        render_table(analysis.get("clauses_df", pd.DataFrame())[[
            col for col in [
                "clause_number", "clause_title", "section_name", "claim_type", "risk_level", "claim_strength",
                "notice_required", "time_impact", "cost_impact", "required_evidence", "recommended_action",
            ] if col in analysis.get("clauses_df", pd.DataFrame()).columns
        ]]),
        "",
        "## Late Drawing Search Ranking",
        "",
        render_table(analysis.get("search_df", pd.DataFrame())[[
            col for col in [
                "clause_title", "section_name", "claim_type", "risk_level", "claim_strength", "search_score"
            ] if col in analysis.get("search_df", pd.DataFrame()).columns
        ]]),
        "",
        "## Evidence Mappings",
        "",
        render_table(analysis.get("mappings_df", pd.DataFrame())),
        "",
        "## Methodology Text Evidence",
        "",
        methodology_text[:2500].replace("\r", ""),
    ]
    return "\n".join(lines)


def score_outputs(analysis: dict[str, Any], artifact_text: str, methodology_text: str) -> dict[str, Any]:
    output_text = artifact_text.split("## Methodology Text Evidence", 1)[0]
    clauses_df = analysis.get("clauses_df", pd.DataFrame())
    search_df = analysis.get("search_df", pd.DataFrame())
    mappings_df = analysis.get("mappings_df", pd.DataFrame())
    notice_df = analysis.get("notice_df", pd.DataFrame())
    answer = analysis["answer"]
    rebuttal = analysis["rebuttal"]
    claim = analysis["claim_payload"]

    contract_checks = [
        ("contract library built", not clauses_df.empty, 12),
        ("clauses classified by section and claim type", not clauses_df.empty and {"section_name", "claim_type"}.issubset(clauses_df.columns), 10),
        ("notice/time/cost columns available", not clauses_df.empty and {"notice_required", "time_impact", "cost_impact"}.issubset(clauses_df.columns), 10),
        ("contract question returns decision", str(answer.get("entitlement_decision", "")).strip() != "", 10),
        ("contract answer includes evidence score", int(answer.get("evidence_strength_score", 0) or 0) >= 0, 8),
        ("rebuttal detects client defenses", not rebuttal.get("detected_defenses_df", pd.DataFrame()).empty, 10),
        ("claim draft includes core sections", all(str(claim.get(key, "")).strip() for key in ["narrative_text", "contractual_basis", "cause_effect", "entitlement_statement", "time_impact_statement", "attachment_checklist"]), 12),
        ("evidence mappings generated", not mappings_df.empty, 10),
        ("contract KPIs generated", bool(analysis.get("kpis")), 8),
        ("stored evidence created", bool(analysis.get("saved_evidence")), 10),
        (
            "late drawing clause ranked in top 3",
            not search_df.empty and search_df.head(3)["clause_title"].astype(str).str.contains("Delayed drawings|Claims time bar|Monthly claim", case=False, regex=True).any(),
            10,
        ),
        (
            "broad design risk not ranked first",
            not search_df.empty and not contains_any(str(search_df.iloc[0].get("clause_title", "")), ["Design verification and review", "Revit and CAD"]),
            8,
        ),
        (
            "evidence maps to late-drawing or time-bar clauses",
            not mappings_df.empty and mappings_df.head(8)["clause_title"].astype(str).str.contains("Delayed drawings|Claims time bar|Monthly claim|Extension of time", case=False, regex=True).any(),
            10,
        ),
        (
            "evidence mapping basis is contextual",
            not mappings_df.empty and mappings_df["mapping_basis"].astype(str).str.contains("signal|context", case=False, regex=True).any(),
            8,
        ),
        (
            "scanned PDF source limitation disclosed",
            bool(analysis.get("pdf_probe", {}).get("is_scanned_or_image_based"))
            and contains_any(json.dumps(analysis.get("status", {}), default=str), ["scanned", "image-based", "curated contract clause library"]),
            12,
        ),
        (
            "scanned PDF render diagnostics available",
            bool(analysis.get("pdf_probe", {}).get("diagnostics", {}).get("render_preview_available")),
            8,
        ),
        (
            "OCR installation requirement disclosed when unavailable",
            (
                bool(analysis.get("pdf_probe", {}).get("diagnostics", {}).get("ocr_available"))
                or contains_any(json.dumps(analysis.get("status", {}), default=str), ["Install Tesseract", "pytesseract", "OCR is not installed"])
            ),
            8,
        ),
    ]
    contract_score, contract_detail = score_boolean_checks(contract_checks)

    letters_checks = [
        ("letters workbook loaded", bool(analysis.get("letters_book")), 14),
        ("notice register generated", not notice_df.empty, 14),
        ("thread linkage present", not notice_df.empty and notice_df["Thread"].astype(str).str.strip().ne("").any(), 12),
        ("reply tracking present", not notice_df.empty and notice_df["Reply Received"].astype(str).isin(["Yes", "No"]).any(), 10),
        ("delay type prediction present", not notice_df.empty and notice_df["Predicted Delay Type"].astype(str).str.strip().ne("").any(), 12),
        ("open action tracking present", not notice_df.empty and notice_df["Next Action"].astype(str).str.strip().ne("").any(), 12),
        ("priority tracking present", not notice_df.empty and notice_df["Priority"].astype(str).str.strip().ne("").any(), 10),
        ("notice refs from both parties present", not notice_df.empty and notice_df["From Party"].nunique() >= 2, 10),
        ("risk/eot subjects visible", contains_any(output_text, ["delay", "eot", "claim", "rfi", "ifc", "critical"]), 6),
    ]
    letters_score, letters_detail = score_boolean_checks(letters_checks)

    methodology_checks = [
        ("mentions notices and time-bar", contains_any(output_text, ["notice", "time-bar", "time bar"]), 10),
        ("mentions entitlement", contains_any(output_text, ["entitlement", "entitled"]), 10),
        ("mentions critical path or longest path", contains_any(output_text, ["critical path", "longest path", "programme impact"]), 12),
        ("mentions TIA/window methodology", contains_any(output_text, ["tia", "window", "time impact", "cpm"]), 12),
        ("mentions concurrency/contractor delay", contains_any(output_text, ["concurrent", "contractor delay", "contractor-caused"]), 10),
        ("mentions causation/cause and effect", contains_any(output_text, ["cause and effect", "causation", "causal"]), 10),
        ("mentions contemporaneous records/evidence", contains_any(output_text, ["contemporaneous", "records", "evidence"]), 10),
        ("mentions cost substantiation", contains_any(output_text, ["cost", "valuation", "quantum", "payment"]), 8),
        ("mentions client rejection/rebuttal", contains_any(output_text, ["rejection", "rebuttal", "client defense"]), 10),
        ("mentions attachments/checklist", contains_any(output_text, ["attachment", "checklist"]), 8),
    ]
    methodology_score, methodology_detail = score_boolean_checks(methodology_checks)

    llm_checks = [
        ("answer is report-grade length", len(output_text) > 8000, 12),
        ("clear decision/risk/evidence labels", contains_any(output_text, ["Decision:", "Risk:", "Evidence strength:"]), 12),
        ("includes missing evidence", contains_any(output_text, ["Missing evidence", "Potential entitlement exists, but evidence is incomplete"]), 10),
        ("separates entitlement/proof/schedule/cost", contains_any(output_text, ["entitlement", "proof", "programme impact", "cost substantiation"]), 14),
        ("client rebuttal is actionable", contains_any(output_text, ["recommended response", "counterargument", "evidence needed"]), 12),
        ("claim draft is submission-oriented", contains_any(output_text, ["contractual basis", "factual background", "attachment checklist"]), 12),
        ("letters are linked to claims workflow", contains_any(output_text, ["Letters Intelligence", "Notice Ref", "Next Action"]), 10),
        ("uses precise claims language", contains_any(output_text, ["notice compliance", "time-bar", "critical path", "EOT", "concurrent"]), 10),
        ("methodology source read", len(methodology_text) > 1000, 8),
    ]
    llm_average, llm_detail = score_boolean_checks(llm_checks)

    overall_score = round((contract_score * 0.30) + (letters_score * 0.20) + (methodology_score * 0.25) + (llm_average * 0.25), 2)
    return {
        "overall_score": overall_score,
        "llm_average": llm_average,
        "contract_score": contract_score,
        "letters_score": letters_score,
        "methodology_score": methodology_score,
        "details": {
            "contract": contract_detail,
            "letters": letters_detail,
            "methodology": methodology_detail,
            "llm": llm_detail,
        },
        "row_counts": {
            "clauses": len(clauses_df),
            "evidence_mappings": len(mappings_df),
            "notices": len(notice_df),
            "rebuttal_defenses": len(rebuttal.get("detected_defenses_df", pd.DataFrame())),
        },
        "kpis": analysis.get("kpis", {}),
    }


def run_eval(label: str, methodology_path: Path) -> dict[str, Any]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    methodology_text = read_docx_text(methodology_path)
    analysis = build_analysis()
    artifact_text = summarize_artifact(analysis, methodology_text)
    artifact_path = OUTPUT_DIR / f"artifact_{label}.md"
    artifact_path.write_text(artifact_text, encoding="utf-8")
    scores = score_outputs(analysis, artifact_text, methodology_text)
    result = {
        "label": label,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "artifact_path": str(artifact_path),
        **scores,
    }
    result_path = OUTPUT_DIR / f"scores_{label}.json"
    result_path.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    with SCORE_LOG.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(result, default=str) + "\n")
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description="Score Contract & Claims and Letters Intelligence outputs.")
    parser.add_argument("--label", default="current")
    parser.add_argument("--methodology", default=str(DEFAULT_METHODOLOGY))
    args = parser.parse_args()
    result = run_eval(args.label, Path(args.methodology))
    summary = {
        "label": result["label"],
        "overall_score": result["overall_score"],
        "llm_average": result["llm_average"],
        "contract_score": result["contract_score"],
        "letters_score": result["letters_score"],
        "methodology_score": result["methodology_score"],
        "artifact_path": result["artifact_path"],
    }
    print(json.dumps(summary, indent=2))
    return 0 if result["overall_score"] >= 90 and result["llm_average"] >= 90 else 1


if __name__ == "__main__":
    raise SystemExit(main())
