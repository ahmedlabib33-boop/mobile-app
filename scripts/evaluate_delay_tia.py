from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
sys.path.insert(0, str(SRC))

from construction_system.steel_delay_tia import (  # noqa: E402
    SteelTiaSettings,
    build_requirement_df_from_client_supply_sheet,
    run_steel_delay_tia_analysis,
)


PROJECTS_DIR = ROOT / "projects"
PROJECT_DIRS = [path for path in PROJECTS_DIR.iterdir() if path.is_dir() and not path.name.startswith("_")]
DEFAULT_PROJECT_DIR = max(PROJECT_DIRS, key=lambda path: sum(file.stat().st_size for file in path.rglob("*") if file.is_file()), default=PROJECTS_DIR / "_PROJECT_TEMPLATE")
DEFAULT_METHODOLOGY = DEFAULT_PROJECT_DIR / "delay_analysis" / "methodology" / "TIA methodology.docx"
TEMPLATE_DIR = DEFAULT_PROJECT_DIR / "delay_analysis" / "steel_delay_tia_templates"
APP_IMPORT_DIR = DEFAULT_PROJECT_DIR / "data" / "import_templates"
OUTPUT_DIR = ROOT / "generated_outputs" / "delay_tia_eval"
SCORE_LOG = OUTPUT_DIR / "score_log.jsonl"


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_csv(name: str) -> pd.DataFrame:
    path = TEMPLATE_DIR / name
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path).fillna("")


def load_app_csv(name: str) -> pd.DataFrame:
    path = APP_IMPORT_DIR / name
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path).fillna("")


def load_first_existing_csv(paths: list[Path]) -> pd.DataFrame:
    for path in paths:
        if path.exists():
            return pd.read_csv(path).fillna("")
    return pd.DataFrame()


def dataframe_records(df: pd.DataFrame, limit: int = 5) -> list[dict[str, Any]]:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return []
    out = df.head(limit).copy()
    for col in out.columns:
        out[col] = out[col].astype(object).where(out[col].notna(), "")
    return out.to_dict(orient="records")


def render_table(frame: pd.DataFrame, limit: int = 8) -> str:
    if frame.empty:
        return "_No rows generated._"
    sample = frame.head(limit).copy()
    for col in sample.columns:
        sample[col] = sample[col].astype(str).str.replace("\r", " ", regex=False).str.replace("\n", " ", regex=False)
    headers = [str(col) for col in sample.columns]
    rows = [[str(value) for value in row] for row in sample.values.tolist()]
    lines = [" | ".join(headers), " | ".join(["---"] * len(headers))]
    lines.extend(" | ".join(row) for row in rows)
    return "\n".join(lines)


def contains_any(text: str, terms: list[str]) -> bool:
    lower = text.lower()
    return any(term.lower() in lower for term in terms)


def pct(value: float, total: float) -> float:
    if total <= 0:
        return 0.0
    return round(max(0.0, min(100.0, value / total * 100.0)), 2)


def score_boolean_checks(checks: list[tuple[str, bool, float]]) -> tuple[float, list[dict[str, Any]]]:
    earned = sum(weight for _, ok, weight in checks if ok)
    total = sum(weight for _, _, weight in checks)
    details = [{"criterion": name, "passed": ok, "weight": weight} for name, ok, weight in checks]
    return pct(earned, total), details


def build_analysis() -> tuple[dict[str, pd.DataFrame | dict[str, Any]], dict[str, pd.DataFrame]]:
    master_df = load_csv("02- master_activity_steel_analysis.csv")
    employer_df = load_first_existing_csv(
        [
            TEMPLATE_DIR / "03- employer_steel_supply_at_site.csv",
            TEMPLATE_DIR / "03- employer_steel_supply.csv",
        ]
    )
    p6_df = load_csv("04- p6_activity_export.csv")
    relationship_df = load_csv("05- relationship_file.csv")
    contract_df = load_csv("06- contract_library.csv")
    rfi_df = load_first_existing_csv(
        [
            TEMPLATE_DIR / "09- rfi_status.csv",
        ]
    )
    delay_events_df = load_app_csv("delay_events.csv")
    requirement_df = build_requirement_df_from_client_supply_sheet(master_df)
    analysis = run_steel_delay_tia_analysis(
        p6_df=p6_df,
        steel_df=employer_df,
        requirement_df=requirement_df,
        relationship_df=relationship_df,
        contract_library_df=contract_df,
        delay_events_df=delay_events_df,
        settings=SteelTiaSettings(
            usability_lag_days=2,
            near_critical_float_threshold=10,
            data_date=pd.Timestamp.today().normalize(),
        ),
    )
    source_frames = {
        "master_df": master_df,
        "employer_df": employer_df,
        "p6_df": p6_df,
        "relationship_df": relationship_df,
        "contract_df": contract_df,
        "rfi_df": rfi_df,
        "delay_events_df": delay_events_df,
        "requirement_df": requirement_df,
    }
    return analysis, source_frames


def summarize_artifact(analysis: dict[str, Any], methodology_text: str) -> str:
    frames = {
        "Affected Candidates": analysis.get("candidates_df", pd.DataFrame()),
        "Fragnet Recommendation": analysis.get("fragnet_df", pd.DataFrame()),
        "Contractual Assessment": analysis.get("assessment_df", pd.DataFrame()),
        "Professional Narrative": analysis.get("narrative_df", pd.DataFrame()),
        "Data Quality": analysis.get("data_quality_df", pd.DataFrame()),
    }
    lines = [
        "# Delay TIA Eval Artifact",
        "",
        f"Generated: {datetime.now(timezone.utc).isoformat()}",
        "",
        "## Methodology Anchors Checked",
        "",
        "- statused schedule windows / update impact analysis",
        "- critical path and float entitlement",
        "- fragnet event, start, recovery, and affected activity sequence",
        "- blindsight handling for multi-window events",
        "- start-of-work impact and duration impact logic",
        "- finish-to-finish logic and lag values",
        "- global vs stepped insertion sequence",
        "- embedded contractor-caused delay and concurrency",
        "- traceable assumptions and final summary",
        "",
        "## Engine Output Samples",
    ]
    for name, frame in frames.items():
        lines.extend([f"### {name}", ""])
        if isinstance(frame, pd.DataFrame) and not frame.empty:
            lines.append(render_table(frame))
        else:
            lines.append("_No rows generated._")
        lines.append("")
    lines.extend(
        [
            "## Methodology Text Evidence",
            "",
            methodology_text[:2500].replace("\r", ""),
        ]
    )
    return "\n".join(lines)


def score_outputs(analysis: dict[str, Any], source_frames: dict[str, pd.DataFrame], artifact_text: str, methodology_text: str) -> dict[str, Any]:
    output_text = artifact_text.split("## Methodology Text Evidence", 1)[0]
    scored_text = output_text.split("## Engine Output Samples", 1)[-1]
    candidates_df = analysis.get("candidates_df", pd.DataFrame())
    fragnet_df = analysis.get("fragnet_df", pd.DataFrame())
    assessment_df = analysis.get("assessment_df", pd.DataFrame())
    narrative_df = analysis.get("narrative_df", pd.DataFrame())
    data_quality_df = analysis.get("data_quality_df", pd.DataFrame())
    kpis = analysis.get("kpis", {}) if isinstance(analysis.get("kpis", {}), dict) else {}

    methodology_checks = [
        ("mentions statused schedule windows", contains_any(scored_text, ["statused", "window", "update impact"]), 10),
        ("mentions critical path and float entitlement", contains_any(scored_text, ["critical path", "float"]), 12),
        ("mentions fragnet start and recovery point", contains_any(scored_text, ["fragment start", "fragment finish", "recovery"]), 12),
        ("mentions first affected activity", contains_any(scored_text, ["first affected", "affected activity"]), 10),
        ("mentions causal path or downstream effect", contains_any(scored_text, ["causal", "downstream", "successor"]), 10),
        ("mentions blindsight or prospective knowledge", contains_any(scored_text, ["blindsight", "anticipated", "known at the start"]), 10),
        ("mentions FF logic or lag values", contains_any(scored_text, ["finish-to-finish", "ff ", "lag"]), 10),
        ("mentions global or stepped insertion", contains_any(scored_text, ["global insertion", "stepped insertion", "chronological"]), 8),
        ("mentions contractor-caused embedded delay", contains_any(scored_text, ["contractor-caused", "contractor caused", "embedded"]), 8),
        ("mentions concurrency and compensability distinction", contains_any(scored_text, ["concurrency", "compensable", "compensation"]), 10),
    ]
    methodology_score, methodology_detail = score_boolean_checks(methodology_checks)

    table_checks = [
        ("candidates generated", isinstance(candidates_df, pd.DataFrame) and not candidates_df.empty, 10),
        ("fragnet rows generated", isinstance(fragnet_df, pd.DataFrame) and not fragnet_df.empty, 10),
        ("assessment rows generated", isinstance(assessment_df, pd.DataFrame) and not assessment_df.empty, 10),
        ("narrative rows generated", isinstance(narrative_df, pd.DataFrame) and not narrative_df.empty, 8),
        ("candidate score present", isinstance(candidates_df, pd.DataFrame) and "TIA Candidate Score" in candidates_df.columns, 8),
        ("criticality present", isinstance(candidates_df, pd.DataFrame) and {"Critical", "Longest Path", "Total Float"}.issubset(candidates_df.columns), 10),
        ("fragnet predecessor successor present", isinstance(fragnet_df, pd.DataFrame) and {"Last completed / available predecessor", "Insert Fragment Before"}.issubset(fragnet_df.columns), 10),
        ("fragment dates and duration present", isinstance(fragnet_df, pd.DataFrame) and {"Fragment Start", "Fragment Finish", "Fragment Duration"}.issubset(fragnet_df.columns), 10),
        ("contract support present", isinstance(assessment_df, pd.DataFrame) and {"Notice Reference", "Final Assessment"}.intersection(assessment_df.columns), 8),
        ("data quality surfaced", isinstance(data_quality_df, pd.DataFrame), 6),
        ("kpis populated", bool(kpis), 10),
    ]
    table_score, table_detail = score_boolean_checks(table_checks)

    llm_checks = [
        ("narrative is report-grade length", len(scored_text) > 7000, 12),
        ("clear uncertainty language", contains_any(scored_text, ["subject to", "not recorded", "not enough data", "pending", "must be"]), 12),
        ("traceable evidence references", contains_any(scored_text, ["uploaded", "register", "p6", "contract", "relationship"]), 14),
        ("executive conclusion possible", contains_any(scored_text, ["final professional opinion", "executive", "conclusion", "recommended"]), 12),
        ("methodology avoids simple sum", contains_any(scored_text, ["not simple addition", "conservative", "concurrency", "double counting"]), 12),
        ("explains entitlement limits", contains_any(scored_text, ["entitlement", "time extension", "compensable", "notice"]), 14),
        ("supports follow-up action", contains_any(scored_text, ["required evidence", "practical action", "action", "mitigation"]), 12),
        ("uses precise schedule terms", contains_any(scored_text, ["critical", "longest path", "float", "predecessor", "successor", "fragnet"]), 12),
    ]
    llm_average, llm_detail = score_boolean_checks(llm_checks)

    source_checks = [
        ("methodology doc read", len(methodology_text) > 1000, 8),
        ("master data loaded", not source_frames["master_df"].empty, 8),
        ("employer steel data loaded", not source_frames["employer_df"].empty, 8),
        ("p6 activity data loaded", not source_frames["p6_df"].empty, 8),
        ("relationship data loaded", not source_frames["relationship_df"].empty, 8),
        ("contract library loaded", not source_frames["contract_df"].empty, 8),
        ("rfi support data loaded", not source_frames["rfi_df"].empty, 6),
        ("quantity basis valid", bool(kpis.get("Quantity Basis Valid")), 10),
        ("strong candidates found", int(kpis.get("Number of Strong TIA Candidates", 0) or 0) > 0, 10),
        ("critical affected activities found", int(kpis.get("Critical Activities Affected", 0) or 0) > 0, 8),
        ("employer risk events found", int(kpis.get("Employer Risk Events", 0) or 0) > 0, 8),
        ("stock-out events found", int(kpis.get("Number of Stock-Out Events", 0) or 0) > 0, 10),
    ]
    data_score, data_detail = score_boolean_checks(source_checks)

    overall_score = round((methodology_score * 0.35) + (table_score * 0.25) + (llm_average * 0.25) + (data_score * 0.15), 2)
    return {
        "overall_score": overall_score,
        "llm_average": llm_average,
        "methodology_score": methodology_score,
        "table_score": table_score,
        "data_score": data_score,
        "details": {
            "methodology": methodology_detail,
            "tables": table_detail,
            "llm": llm_detail,
            "data": data_detail,
        },
        "kpis": kpis,
        "row_counts": {
            "candidates_df": len(candidates_df) if isinstance(candidates_df, pd.DataFrame) else 0,
            "fragnet_df": len(fragnet_df) if isinstance(fragnet_df, pd.DataFrame) else 0,
            "assessment_df": len(assessment_df) if isinstance(assessment_df, pd.DataFrame) else 0,
            "narrative_df": len(narrative_df) if isinstance(narrative_df, pd.DataFrame) else 0,
        },
    }


def run_eval(label: str, methodology_path: Path) -> dict[str, Any]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    methodology_text = read_docx_text(methodology_path)
    analysis, source_frames = build_analysis()
    artifact_text = summarize_artifact(analysis, methodology_text)
    artifact_path = OUTPUT_DIR / f"artifact_{label}.md"
    artifact_path.write_text(artifact_text, encoding="utf-8")
    scores = score_outputs(analysis, source_frames, artifact_text, methodology_text)
    result = {
        "label": label,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "artifact_path": str(artifact_path),
        **scores,
    }
    result_path = OUTPUT_DIR / f"scores_{label}.json"
    result_path.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    with SCORE_LOG.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(result, default=str) + "\n")
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description="Score Delay TIA output against the methodology document.")
    parser.add_argument("--label", default="current", help="Label used in generated artifact filenames.")
    parser.add_argument("--methodology", default=str(DEFAULT_METHODOLOGY), help="Path to TIA methodology DOCX.")
    args = parser.parse_args()
    result = run_eval(args.label, Path(args.methodology))
    summary = {
        "label": result["label"],
        "overall_score": result["overall_score"],
        "llm_average": result["llm_average"],
        "methodology_score": result["methodology_score"],
        "table_score": result["table_score"],
        "data_score": result["data_score"],
        "artifact_path": result["artifact_path"],
    }
    print(json.dumps(summary, indent=2))
    return 0 if result["overall_score"] >= 90 and result["llm_average"] >= 90 else 1


if __name__ == "__main__":
    raise SystemExit(main())
